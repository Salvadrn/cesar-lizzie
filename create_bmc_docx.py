#!/usr/bin/env python3
"""Generate a professional Word document for the NeuroNav Business Model Canvas."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# --- Styles ---
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)

# --- Title ---
title = doc.add_heading('NeuroNav', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Business Model Canvas')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
run.bold = True

subtitle2 = doc.add_paragraph()
subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = subtitle2.add_run('The Neuronavs  |  Tech4Good 2026')
run2.font.size = Pt(12)
run2.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

subtitle3 = doc.add_paragraph()
subtitle3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = subtitle3.add_run('Adaptive Daily Living Assistant for Adults with Cognitive Disabilities')
run3.font.size = Pt(11)
run3.font.italic = True
run3.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()

# --- Helper function ---
def add_section(title_text, color, items, description=""):
    heading = doc.add_heading(title_text, level=1)
    for run in heading.runs:
        run.font.color.rgb = color

    if description:
        desc_para = doc.add_paragraph()
        desc_run = desc_para.add_run(description)
        desc_run.font.italic = True
        desc_run.font.size = Pt(10)
        desc_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    for item_title, item_desc in items:
        para = doc.add_paragraph(style='List Bullet')
        bold_run = para.add_run(item_title)
        bold_run.bold = True
        if item_desc:
            para.add_run(f' - {item_desc}')

    doc.add_paragraph()

# Colors
GOLD = RGBColor(0xB8, 0x86, 0x0B)
RED = RGBColor(0xC0, 0x39, 0x2B)
TEAL = RGBColor(0x1A, 0x7A, 0x6D)
BLUE = RGBColor(0x1A, 0x56, 0xDB)

# --- 1. Key Partners ---
add_section('1. Key Partners', GOLD,
    [
        ('Hospitals & Rehabilitation Centers', 'Healthcare facilities that treat patients with cognitive disabilities and can recommend the app as part of their therapy program.'),
        ('Disability Organizations & Nonprofits', 'Foundations and support groups that serve adults with cognitive disabilities and connect us with potential users.'),
        ('Apple', 'Technology partner providing the iOS, watchOS, and WidgetKit ecosystem that powers our multi-device experience.'),
        ('Supabase', 'Cloud infrastructure provider for our database, authentication, and real-time data synchronization services.'),
        ('Universities & Research Centers', 'Academic institutions that help with clinical validation, user testing, and evidence-based design.'),
        ('Occupational Therapists', 'Healthcare professionals who recommend and guide the use of assistive technology with their patients.'),
        ('Government Health Agencies', 'Public institutions that fund assistive technology programs and support inclusive initiatives.'),
    ],
    'Who are our key partners and suppliers?'
)

# --- 2. Key Activities ---
add_section('2. Key Activities', GOLD,
    [
        ('App Development & Maintenance', 'Building and updating the iOS, watchOS, and Flutter applications with accessible interfaces.'),
        ('Adaptive UX Research', 'Conducting research to create interfaces that automatically adjust to each user\'s cognitive ability level.'),
        ('User Testing with Real Patients', 'Running usability tests with adults who have cognitive disabilities to validate our design choices.'),
        ('Caregiver Onboarding & Training', 'Creating resources and guides so caregivers can effectively set up and manage routines for their patients.'),
        ('Data Analysis & Engine Improvement', 'Analyzing routine completion data to continuously improve our 5-level adaptive complexity engine.'),
        ('Partnership Development', 'Building relationships with healthcare organizations, disability groups, and technology partners.'),
        ('Routine Template Creation', 'Designing pre-built, step-by-step routine templates for common daily activities (cooking, hygiene, medication, etc.).'),
    ],
    'What key activities does our value proposition require?'
)

# --- 3. Key Resources ---
add_section('3. Key Resources', GOLD,
    [
        ('Development Team', 'Engineers with expertise in Swift, Flutter, NestJS, and accessible design principles.'),
        ('Supabase Cloud Infrastructure', 'PostgreSQL database with Row-Level Security, authentication, and real-time subscriptions.'),
        ('Apple Developer Program', 'Membership required for App Store distribution and access to Apple frameworks (WidgetKit, App Intents, HealthKit).'),
        ('Adaptive Engine Algorithm', 'Proprietary heuristic engine that adjusts task complexity across 5 levels based on user performance.'),
        ('Pre-built Routine Library', 'A collection of step-by-step guides for 9 categories of daily activities.'),
        ('Apple Watch Hardware', 'Wearable device for haptic feedback, wrist notifications, and health monitoring integration.'),
        ('Clinical Advisor Network', 'Medical and therapeutic professionals who guide our product decisions.'),
    ],
    'What key resources does our value proposition require?'
)

# --- 4. Value Proposition ---
add_section('4. Value Proposition', RED,
    [
        ('Adaptive Task Guidance', 'The app automatically adjusts task complexity based on user performance. If a user struggles, steps become simpler. If they succeed, steps become more challenging.'),
        ('Step-by-Step Guided Routines', 'Clear, visual instructions for daily activities like cooking, hygiene, medication management, laundry, and transportation.'),
        ('Real-Time Stall Detection', 'The app detects when a user is stuck on a step and automatically offers help through voice prompts and simplified instructions.'),
        ('Caregiver-Patient Connection', 'Caregivers can create routines, monitor progress, and receive alerts remotely, building a support network around each patient.'),
        ('Multi-Sensory Feedback', 'Combines voice guidance (text-to-speech), haptic feedback (Apple Watch vibrations), and visual cues for maximum accessibility.'),
        ('Safety Features', 'Geofencing with safety zones and emergency contacts ensures user safety during transit and daily activities.'),
        ('Medication & Appointment Management', 'Built-in tracking for medications and medical appointments with reminder notifications.'),
        ('Privacy-First Design', 'Row-Level Security in the database ensures each user can only access their own data.'),
    ],
    'What value do we deliver to the customer?'
)

# --- 5. Customer Relationships ---
add_section('5. Customer Relationships', TEAL,
    [
        ('Patient-Caregiver Bond', 'Caregivers set up and customize routines; patients receive step-by-step guidance and support during execution.'),
        ('Family Linking', 'Family members can monitor their loved one\'s progress and routine completion without being able to modify data.'),
        ('Personalized Adaptation', 'The adaptive engine learns from each user\'s performance history and automatically adjusts difficulty levels.'),
        ('Automated Assistance', 'Push notifications, text-to-speech voice guidance, and Apple Watch haptics provide continuous support without requiring human intervention.'),
        ('Trust & Safety', 'All personal data is private and protected. Users and caregivers feel safe knowing their information is secure.'),
        ('Community Building', 'Connecting families and caregivers through shared experiences and mutual support.'),
    ],
    'What type of relationship does each customer segment expect?'
)

# --- 6. Channels ---
add_section('6. Channels', TEAL,
    [
        ('Apple App Store', 'Primary distribution channel for the iOS and watchOS applications.'),
        ('Google Play Store', 'Distribution channel for the Android version (built with Flutter).'),
        ('Hospitals & Clinics', 'Healthcare facilities that recommend the app to patients as part of their treatment plan.'),
        ('Disability Organizations', 'Nonprofits and foundations that distribute information about assistive technology tools.'),
        ('Social Media', 'Instagram and Facebook campaigns targeting caregivers, families, and healthcare professionals.'),
        ('Healthcare Conferences', 'Events and expos where we can demonstrate the product and build partnerships.'),
        ('Web Dashboard', 'Browser-based management tool for caregivers to create and monitor routines from a computer.'),
        ('Word of Mouth', 'Referrals from satisfied caregivers, families, and patients.'),
    ],
    'Through which channels do our customer segments want to be reached?'
)

# --- 7. Customer Segments ---
add_section('7. Customer Segments', TEAL,
    [
        ('Adults with Cognitive Disabilities (Primary Users)', 'People with acquired brain injuries, intellectual disabilities, or neurodegenerative conditions who need support with daily activities.'),
        ('Professional Caregivers', 'Occupational therapists, home care aides, and support workers in group homes who manage routines for their patients.'),
        ('Family Caregivers', 'Family members who provide daily care and want to help their loved ones become more independent.'),
        ('Family Members (Observers)', 'Relatives who want to monitor progress and stay informed without directly managing routines.'),
        ('Healthcare Institutions', 'Hospitals, clinics, and rehabilitation centers that want to integrate assistive technology into their programs.'),
        ('Disability Organizations', 'Nonprofits and foundations that support adults with cognitive disabilities.'),
    ],
    'For whom are we creating value?'
)

# --- 8. Cost Structure ---
add_section('8. Cost Structure', BLUE,
    [
        ('Development Costs', 'Salaries and resources for iOS (Swift), Flutter, and backend (NestJS) developers.'),
        ('Cloud Infrastructure', 'Supabase hosting fees for database, authentication, real-time sync, and storage.'),
        ('Apple Developer Program', '$99/year membership fee for App Store distribution.'),
        ('Google Play Developer', '$25 one-time registration fee for Play Store distribution.'),
        ('User Research & Testing', 'Costs associated with clinical user studies and accessibility testing.'),
        ('Marketing & Outreach', 'Social media advertising, conference attendance, and promotional materials.'),
        ('Design & Accessibility Audits', 'Professional reviews to ensure compliance with accessibility standards (WCAG).'),
        ('Legal & Compliance', 'Data protection compliance (GDPR, local privacy laws).'),
        ('Customer Support', 'Helping caregivers and families with app setup and troubleshooting.'),
    ],
    'What are the most important costs in our business model?'
)

# --- 9. Revenue Streams ---
add_section('9. Revenue Streams', BLUE,
    [
        ('Freemium Model', 'Core features are free to attract users. Premium features are available through paid subscriptions.'),
        ('Basic Plan (Free)', '1 patient profile, up to 3 routines, basic progress tracking.'),
        ('Standard Plan ($4.99/month)', 'Unlimited routines, caregiver linking, safety zones, medication tracking.'),
        ('Professional Plan ($14.99/month)', 'Multi-patient management, analytics dashboard, priority support, institutional features.'),
        ('Institutional Licensing', 'Volume pricing for hospitals, clinics, and care facilities managing multiple patients.'),
        ('Government Grants & Funding', 'Applying for assistive technology grants and disability inclusion programs.'),
        ('Insurance Partnerships', 'Working with insurance companies to cover the app as an assistive technology benefit.'),
    ],
    'For what value are our customers willing to pay?'
)

# --- Tech Stack Table ---
doc.add_heading('Technology Stack', level=1)
table = doc.add_table(rows=9, cols=2, style='Light Shading Accent 1')
table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = table.rows[0].cells
headers[0].text = 'Component'
headers[1].text = 'Technology'

data = [
    ('iOS App', 'Swift 5.9+, SwiftUI, iOS 17+'),
    ('Watch App', 'watchOS 10+, WatchConnectivity'),
    ('Mobile App', 'Flutter (cross-platform)'),
    ('Backend', 'NestJS + TypeORM + PostgreSQL'),
    ('Database', 'Supabase (PostgreSQL + Auth + Real-time)'),
    ('Dashboard', 'Next.js 15 + Tailwind CSS'),
    ('Widgets', 'WidgetKit + Live Activities'),
    ('Voice Assistant', 'Siri (App Intents) + Text-to-Speech'),
]

for i, (comp, tech) in enumerate(data):
    row = table.rows[i + 1].cells
    row[0].text = comp
    row[1].text = tech

doc.add_paragraph()

# --- Competitive Advantage ---
doc.add_heading('Competitive Advantage', level=1)

advantages = [
    ('Adaptive Complexity Engine', 'No other app on the market adjusts task difficulty in real-time based on user performance. Our 5-level system ensures every user gets the right amount of support.'),
    ('Multi-Sensory Feedback', 'We combine voice guidance, haptic vibrations, and visual cues to reach users through multiple channels, maximizing comprehension and engagement.'),
    ('Complete Caregiver Ecosystem', 'A full platform connecting patients, caregivers, and families with role-based access and remote monitoring capabilities.'),
    ('Apple Watch Integration', 'Wrist-based step guidance with haptic prompts and health monitoring creates a discreet, always-available support system.'),
    ('Privacy-First Architecture', 'Row-Level Security in our database ensures strict data isolation. Each user can only access their own information.'),
]

for i, (title_text, desc) in enumerate(advantages, 1):
    para = doc.add_paragraph()
    num_run = para.add_run(f'{i}. ')
    num_run.bold = True
    title_run = para.add_run(title_text)
    title_run.bold = True
    para.add_run(f' - {desc}')

doc.add_paragraph()

# --- Footer ---
footer_para = doc.add_paragraph()
footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_run = footer_para.add_run('NeuroNav - Empowering Independence Through Adaptive Technology')
footer_run.font.italic = True
footer_run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)
footer_run.font.size = Pt(12)

footer2 = doc.add_paragraph()
footer2.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer2_run = footer2.add_run('Tech4Good 2026')
footer2_run.font.size = Pt(10)
footer2_run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

# Save
output_path = '/Users/salvador/Desktop/Apps Kirks/Proyecto2/NeuroNav_Business_Model_Canvas.docx'
doc.save(output_path)
print(f'Document saved to: {output_path}')
