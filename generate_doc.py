from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# --- Styles ---
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# --- PORTADA ---
for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Soluciones Tecnológicas Bien Pensadas\npara la Salud e Inclusión')
run.bold = True
run.font.size = Pt(26)
run.font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Proyecto Final — Etapa 1: Investigación')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run('Categoría: Salud + Inclusión\nFebrero 2026')
run.font.size = Pt(12)

doc.add_page_break()

# --- INTRODUCCIÓN ---
h = doc.add_heading('Introducción', level=1)
h.runs[0].font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)

doc.add_paragraph(
    'La tecnología tiene el poder de transformar la vida de personas con discapacidades cognitivas '
    'y adultos mayores que enfrentan retos en sus actividades diarias. En este documento se analizan '
    '5 soluciones tecnológicas existentes que abordan problemáticas reales en el ámbito de la salud '
    'y la inclusión, evaluando por qué están bien pensadas según los criterios vistos en clase: '
    'diseño centrado en el usuario, accesibilidad, impacto social, viabilidad técnica y sostenibilidad.'
)

doc.add_paragraph()

# --- SOLUCIONES ---
solutions = [
    {
        'name': '1. Medisafe — Recordatorio de Medicamentos',
        'url': 'medisafe.com',
        'what': (
            'Medisafe es una plataforma integral de gestión de medicamentos que ayuda a los usuarios '
            'a seguir horarios complejos de medicación mediante alertas personalizadas, advertencias '
            'de interacciones medicamentosas peligrosas, y seguimiento de mediciones de salud. '
            'Su función "Medfriend" notifica a un cuidador designado cuando se olvida una dosis.'
        ),
        'why': [
            'Resuelve un problema real y crítico: la adherencia a medicamentos, que causa 125,000 muertes anuales solo en EE.UU.',
            'Ha prevenido más de 200,000 interacciones peligrosas entre medicamentos, actuando como red de seguridad.',
            'Disponible en 16+ idiomas, incluyendo español, con soporte para Apple Watch — reduciendo la fricción para adultos mayores.',
            'La función "Medfriend" crea una capa de responsabilidad compartida entre paciente y cuidador.',
            'En 2025 introdujo VIA (Voice Intelligent Agent), permitiendo interacción por voz para personas que luchan con interfaces táctiles.',
            'Cumple con HIPAA y GDPR con encriptación de 256 bits, protegiendo datos médicos sensibles.',
        ],
        'tech': 'Push notifications, SMS, escaneo de código de barras, integración con Apple Health, agente de voz con IA generativa, iOS/Android/Apple Watch.',
    },
    {
        'name': '2. ElliQ — Robot Acompañante con IA para Adultos Mayores',
        'url': 'elliq.com',
        'what': (
            'ElliQ es un robot compañero proactivo desarrollado por Intuition Robotics, diseñado para '
            'adultos mayores que viven solos. Inicia conversaciones sin esperar comandos, sugiere actividades '
            '(ejercicio, mindfulness, viajes virtuales a museos), proporciona recordatorios de medicamentos, '
            'y conecta a los usuarios con su familia a través de videollamadas.'
        ),
        'why': [
            'Diseño proactivo: a diferencia de asistentes que esperan comandos, ElliQ inicia interacción — esencial para personas con deterioro cognitivo que no recuerdan pedir ayuda.',
            'Estudio del estado de Nueva York demostró 95% de reducción en soledad auto-reportada.',
            'Usuarios interactúan con ElliQ más de 30 veces al día, 6 días a la semana — indicando engagement real.',
            'Usa interacción multimodal (voz, movimiento corporal, luces, texto, imágenes) siendo accesible sin importar el modo de comunicación preferido.',
            'Su "Relationship Orchestration Engine" construye una relación persistente recordando conversaciones, estados emocionales y preferencias.',
            'La solución para cuidadores ($9.99/mes) permite monitoreo remoto sin requerir nada del adulto mayor.',
        ],
        'tech': 'Motor de orquestación relacional, LLM con guardarraíles de seguridad, extracción de memorias, salida multimodal, app para cuidadores.',
    },
    {
        'name': '3. Life360 — Monitoreo de Seguridad Familiar',
        'url': 'life360.com',
        'what': (
            'Life360 es una plataforma de seguridad familiar que ofrece ubicación GPS en tiempo real, '
            'alertas de geo-cercas, detección de accidentes automovilísticos, alertas SOS, y análisis '
            'de manejo seguro. Para cuidadores, permite monitorear la ubicación de un ser querido con '
            'deterioro cognitivo y recibir alertas cuando sale de zonas seguras.'
        ),
        'why': [
            'Más de 66 millones de usuarios globalmente — una de las plataformas de seguridad familiar más probadas.',
            'Las geo-cercas atacan directamente el problema de deambulación en Alzheimer: alertas instantáneas cuando el paciente sale de una zona segura.',
            'La detección de accidentes no requiere acción del usuario — crítico para personas con discapacidad cognitiva desorientadas tras un accidente.',
            'Interfaz simple diseñada para que adultos mayores con experiencia tecnológica limitada la puedan usar.',
            'Diseño consciente de privacidad con "Circles": la ubicación solo es visible para miembros aprobados.',
            'Alertas SOS envían ubicación precisa silenciosamente a todos los contactos de emergencia.',
        ],
        'tech': 'GPS híbrido + WiFi + datos celulares, geo-cercas configurables, detección telemática de accidentes, historial de ubicación de 30 días, iOS/Android.',
    },
    {
        'name': '4. Senior Safety App — Emergencias para Demencia',
        'url': 'seniorsafetyapp.com',
        'what': (
            'Senior Safety App es una plataforma de seguridad para adultos mayores con demencia que ofrece '
            'seguimiento GPS en tiempo real, alertas de geo-cercas, detección de caídas, monitoreo de '
            'inactividad, botón SOS, y protección contra estafas y fraudes telefónicos.'
        ),
        'why': [
            'Cada función mapea a un riesgo clínico específico de la demencia: geo-cercas para deambulación, detección de caídas, alertas de inactividad.',
            'El monitoreo de inactividad es único: si el adulto mayor no se mueve por X horas, alerta al cuidador — detecta emergencias silenciosas como derrames cerebrales.',
            'La protección contra fraudes aborda una vulnerabilidad frecuentemente ignorada: personas con deterioro cognitivo son blanco desproporcionado de estafas.',
            'Precio accesible: $4.50/mes o $45/año, haciéndolo viable para familias de todos los niveles de ingreso.',
            'El historial de rutas permite a cuidadores identificar patrones (visitas repetidas a lugares desconocidos pueden indicar confusión).',
            'La función de "solicitar ayuda" envía mensajes de texto Y correos electrónicos con GPS, creando redundancia en comunicación de emergencia.',
        ],
        'tech': 'GPS 24/7, geo-cercas configurables, detección de caídas por acelerómetro, temporizador de inactividad, SOS con llamadas en cascada, iOS/Android.',
    },
    {
        'name': '5. MindMate — Salud Cognitiva Adaptativa',
        'url': 'mindmate-app.com',
        'what': (
            'MindMate es una app de salud integral diseñada como "compañero digital" para personas con '
            'demencia o deterioro cognitivo. Combina juegos cerebrales (velocidad, memoria, resolución '
            'de problemas), rutinas de actividad diaria, recetas nutritivas, programas de ejercicio físico '
            'y artículos educativos sobre envejecimiento saludable.'
        ),
        'why': [
            'Fundada por 3 cuidadores que vivieron personalmente el deterioro de sus seres queridos — diseño impulsado por empatía real.',
            'Desarrollada con guía del Dr. Terry Quinn de la Universidad de Glasgow, basada en investigación revisada por pares.',
            'Estudio publicado en PubMed demostró mejora estadísticamente significativa en rendimiento de memoria (p < .01).',
            'Enfoque holístico: combina entrenamiento cognitivo, nutrición, ejercicio y socialización en un solo régimen diario.',
            'Personalización desde el primer uso: el onboarding pregunta género, edad y condiciones para adaptar la experiencia.',
            'Modo practicante permite a profesionales de salud usar MindMate como herramienta clínica para seguimiento de pacientes.',
        ],
        'tech': 'Juegos cerebrales basados en evidencia, regímenes diarios personalizados, guía nutricional, seguimiento de progreso, modo practicante, iOS/Android.',
    },
]

for sol in solutions:
    h = doc.add_heading(sol['name'], level=2)
    h.runs[0].font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)

    p = doc.add_paragraph()
    run = p.add_run('🌐 ')
    run = p.add_run(sol['url'])
    run.font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)
    run.underline = True

    p = doc.add_paragraph()
    run = p.add_run('¿Qué hace?')
    run.bold = True
    p.add_run('\n' + sol['what'])

    p = doc.add_paragraph()
    run = p.add_run('¿Por qué está bien pensada?')
    run.bold = True

    for reason in sol['why']:
        doc.add_paragraph(reason, style='List Bullet')

    p = doc.add_paragraph()
    run = p.add_run('Tecnología utilizada: ')
    run.bold = True
    p.add_run(sol['tech'])

    doc.add_paragraph()  # spacing

# --- TABLA COMPARATIVA ---
doc.add_page_break()
h = doc.add_heading('Tabla Comparativa', level=1)
h.runs[0].font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)

doc.add_paragraph(
    'La siguiente tabla compara las funcionalidades de las 5 soluciones investigadas '
    'con NeuroNav, nuestra propuesta de solución tecnológica:'
)

headers = ['Funcionalidad', 'Medisafe', 'ElliQ', 'Life360', 'SeniorSafety', 'MindMate', 'NeuroNav']
rows = [
    ['Recordatorios de medicamentos', '✅', '✅', '❌', '❌', '❌', '✅'],
    ['Detección de caídas', '❌', '❌', '✅', '✅', '❌', '✅'],
    ['Rutinas adaptativas', '❌', 'Parcial', '❌', '❌', '✅', '✅'],
    ['Modo perdido', '❌', '❌', 'Parcial', 'Parcial', '❌', '✅'],
    ['Modo simple / accesible', '❌', '✅', '❌', '❌', '❌', '✅'],
    ['Vínculo con cuidador', '✅', '✅', '✅', '✅', 'Parcial', '✅'],
    ['SOS / Emergencia', '❌', '❌', '✅', '✅', '❌', '✅'],
    ['Niveles de complejidad', '❌', '❌', '❌', '❌', 'Parcial', '✅'],
    ['Multiplataforma', '✅', '❌', '✅', '✅', '✅', '✅'],
]

table = doc.add_table(rows=1 + len(rows), cols=len(headers))
table.style = 'Medium Shading 1 Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

for i, h_text in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h_text
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True
            run.font.size = Pt(9)

for r_idx, row in enumerate(rows):
    for c_idx, val in enumerate(row):
        cell = table.rows[r_idx + 1].cells[c_idx]
        cell.text = val
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
            for run in paragraph.runs:
                run.font.size = Pt(9)

# --- CONCLUSIÓN ---
doc.add_paragraph()
h = doc.add_heading('Conclusión', level=1)
h.runs[0].font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)

doc.add_paragraph(
    'Las 5 soluciones analizadas demuestran que una solución tecnológica bien pensada en el ámbito '
    'de la salud e inclusión debe cumplir con principios fundamentales:'
)

conclusions = [
    'Diseño centrado en el usuario: Entender las limitaciones cognitivas, físicas y tecnológicas de la población objetivo.',
    'Accesibilidad universal: Interfaces simplificadas, interacción multimodal (voz, táctil, visual), y personalización.',
    'Impacto medible: Respaldadas por estudios, datos de uso, o evidencia clínica que demuestran resultados reales.',
    'Red de apoyo: Involucran al ecosistema completo (paciente, cuidador, profesional de salud), no solo al usuario final.',
    'Seguridad y privacidad: Protección de datos sensibles de salud con estándares como HIPAA y encriptación.',
]

for c in conclusions:
    doc.add_paragraph(c, style='List Bullet')

doc.add_paragraph(
    '\nNeuroNav busca integrar las mejores prácticas de estas 5 soluciones en una sola plataforma '
    'adaptativa, multiplataforma, y diseñada específicamente para adultos con discapacidades '
    'cognitivas en nuestra comunidad.'
)

# --- Save ---
output_path = '/Users/salvador/Desktop/Investigacion_NeuroNav.docx'
doc.save(output_path)
print(f'Documento guardado en: {output_path}')
